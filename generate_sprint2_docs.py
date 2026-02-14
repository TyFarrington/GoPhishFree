"""
Generate Sprint 2 Requirements and Artifacts as .docx (Word) files
for GoPhishFree - EECS582 Capstone Project - Team 24

Matches the formatting style of the Sprint 1 documents.
Written as sprint planning documents (goals, not completed items).
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUT_DIR  = os.path.join(BASE_DIR, "docs", "Sprint2")
os.makedirs(OUT_DIR, exist_ok=True)

TEAM = "Ty Farrington, Andrew Reyes, Brett Suhr, Nicholas Holmes, Kaleb Howard"
SPRINT_PERIOD = "February 9 - February 22, 2026"


def set_table_style(table):
    """Apply a clean bordered style to a table."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.style.font.size = Pt(10)


def add_header_row(table, headers):
    """Bold the first row of a table as a header."""
    for i, cell in enumerate(table.rows[0].cells):
        cell.text = headers[i]
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)


def add_data_row(table, values):
    """Add a row of data to a table."""
    row = table.add_row()
    for i, val in enumerate(values):
        row.cells[i].text = str(val)
        for p in row.cells[i].paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)
    return row


# ===================================================================
# DOCUMENT 1: Sprint 2 Requirements
# ===================================================================
def build_requirements():
    doc = Document()

    # -- Title --
    doc.add_heading("GoPhishFree - Sprint 2 Requirements", level=1)

    doc.add_paragraph(
        "Project: GoPhishFree - AI-Powered Phishing Detection Chrome Extension"
    )
    doc.add_paragraph(f"Team: {TEAM}")
    doc.add_paragraph(f"Sprint Period: {SPRINT_PERIOD}")

    # -- Overview --
    doc.add_heading("Overview", level=2)
    doc.add_paragraph(
        "Sprint 2 will focus on the core machine learning pipeline: designing, training, "
        "validating, and integrating a unified calibrated phishing detection model into the "
        "Chrome extension built during Sprint 1. The model will use a 64-feature unified schema "
        "covering URL analysis, custom rule signals, DNS checks, Deep Scan page features, "
        "BEC/linkless detection, and attachment metadata - all in a single Random Forest with "
        "isotonic calibration. We also plan to implement a comprehensive trusted domain whitelist "
        "with post-model intelligence to reduce false positives, a free email provider distinction "
        "to prevent whitelist abuse, user-managed custom trusted domains, and an optional cloud AI "
        "enhancement (BYOK) feature to provide a second opinion using features-only payloads. By "
        "the end of this sprint the extension should generate a calibrated risk score (0-100) for "
        "every email opened in Gmail, display contributing risk factors, and optionally query an AI "
        "provider for additional analysis. All core ML processing will remain fully local."
    )

    # -- Sprint 2 Requirements Table --
    doc.add_heading("Sprint 2 Requirements (From Requirements Stack)", level=2)
    doc.add_paragraph(
        "The following requirements are assigned to Sprint 2 in the project requirements "
        "stack. These are the primary goals the team will work toward completing this sprint."
    )

    sprint2_reqs = [
        ("6",  "Develop an initial phishing detection model design", "5", "High"),
        ("7",  "Train the phishing detection model using labeled phishing email datasets", "8", "High"),
        ("8",  "Validate the trained phishing detection model for classification accuracy", "5", "High"),
        ("9",  "Integrate the trained phishing detection model into the extension's analysis pipeline", "8", "High"),
        ("10", "Generate a phishing risk score for the selected email using the integrated model", "5", "High"),
        ("11", "Display the phishing risk score within the report dashboard", "3", "Medium"),
        ("12", "Display the primary reasons contributing to the phishing risk score", "5", "Medium"),
        ("13", "Identify and display suspicious links found in the email", "5", "Medium"),
        ("21", "Display loading indicators while email analysis is in progress", "2", "Medium"),
        ("22", "Display clear error messages when analysis cannot be completed", "3", "Medium"),
    ]

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    add_header_row(table, ["Req ID", "Requirement", "Story Pts", "Priority"])
    for req in sprint2_reqs:
        add_data_row(table, req)
    set_table_style(table)

    # -- Stretch Goals --
    doc.add_heading("Stretch Goals (Ahead-of-Schedule Items)", level=2)
    doc.add_paragraph(
        "If time permits, the team will also begin work on the following items from later "
        "sprints to get ahead of the project timeline."
    )

    stretch = [
        ("Unified 64-feature calibrated model",
         "Req 15 (Sprint 3)",
         "Replace the two-model pipeline (25-feature Tier 1 + 38-feature Deep Scan) with a "
         "single 64-feature unified Random Forest using isotonic calibration. "
         "riskScore = round(100 x calibrated_probability)."),
        ("BEC/Linkless detection features",
         "Req 15 (Sprint 3)",
         "Add 5 new features for business email compromise: FinancialRequestScore, "
         "AuthorityImpersonationScore, PhoneCallbackPattern, ReplyToMismatch, IsLinkless."),
        ("Attachment analysis features",
         "Req 15 (Sprint 3)",
         "Add 5 new features for attachment threats: HasAttachment, AttachmentCount, "
         "RiskyAttachmentExtension, DoubleExtensionFlag, AttachmentNameEntropy."),
        ("Trusted domain whitelist (500+ domains)",
         "Req 16 (Sprint 3)",
         "Implement a comprehensive whitelist of 500+ known-legitimate corporate/organizational "
         "domains. Separate free email providers (gmail.com, outlook.com, etc.) into a distinct "
         "set to prevent whitelist abuse by phishers using free accounts."),
        ("Post-model intelligence layer",
         "Req 16 (Sprint 3)",
         "Add rule-based post-model adjustments: BEC boosts (floor 70-80 for strong financial/"
         "authority signals), trusted domain dampening (cap at 30), newsletter detection (cap "
         "at 45), and free email provider awareness (no dampening for public providers)."),
        ("User-managed custom trusted domains",
         "Req 16 (Sprint 3)",
         "Allow users to add/remove custom trusted domains and block built-in ones via the "
         "popup settings UI. Store in chrome.storage.local with live sync to content script."),
        ("AI Enhancement (Cloud BYOK)",
         "Req 27 (Sprint 4)",
         "Add optional cloud AI second opinion using BYOK (Bring Your Own Key). Support "
         "OpenAI, Anthropic, Google Gemini, Azure OpenAI, and custom endpoints. "
         "Features-only payload (no email body/subject/sender). Automatic gating to "
         "reduce latency/cost. Strict JSON schema validation."),
        ("Synthetic training data augmentation",
         "Req 15 (Sprint 3)",
         "Generate synthetic BEC phishing (2,000), attachment phishing (1,000), legitimate "
         "newsletter (2,000), and transactional email (1,000) samples to improve model "
         "accuracy on underrepresented phishing types."),
        ("Fish award system and collection panel",
         "Req 17 (Sprint 3)",
         "Award themed fish (Friendly Fish, Suspicious Fish, Phishy Puffer, Mega Phish "
         "Shark) when an email scan completes. Animated SVG fish tank dashboard."),
        ("DNS-over-HTTPS validation",
         "Req 27 (Sprint 4)",
         "Implement domain validation using Cloudflare and Google DNS-over-HTTPS, "
         "integrated as model input features."),
    ]

    table2 = doc.add_table(rows=1, cols=3)
    table2.style = "Table Grid"
    add_header_row(table2, ["Item", "Related Req", "Description"])
    for s in stretch:
        add_data_row(table2, s)
    set_table_style(table2)

    # -- Planning Summary --
    doc.add_heading("Sprint 2 Planning Summary", level=2)

    summary = [
        ("Sprint 2 Requirements", "10"),
        ("Sprint 2 Story Points", "49"),
        ("Stretch Goal Items", "10"),
        ("Sprint Duration", "14 days"),
        ("Team Members", "5"),
    ]

    table3 = doc.add_table(rows=1, cols=2)
    table3.style = "Table Grid"
    add_header_row(table3, ["Metric", "Value"])
    for s in summary:
        add_data_row(table3, s)
    set_table_style(table3)

    out = os.path.join(OUT_DIR, "Sprint2_Requirements.docx")
    doc.save(out)
    print(f"  -> {out}")


# ===================================================================
# DOCUMENT 2: Sprint 2 Artifacts
# ===================================================================
def build_artifacts():
    doc = Document()

    # -- Title --
    doc.add_heading("GoPhishFree - Sprint 2 Artifacts", level=1)

    doc.add_paragraph(
        "Project: GoPhishFree - AI-Powered Phishing Detection Chrome Extension"
    )
    doc.add_paragraph(f"Team: {TEAM}")
    doc.add_paragraph(f"Sprint Period: {SPRINT_PERIOD}")

    # -- Overview --
    doc.add_heading("Overview", level=2)
    doc.add_paragraph(
        "This document outlines the artifacts to be produced during Sprint 2 of the GoPhishFree "
        "project. Sprint 2 deliverables will focus on the unified calibrated ML model pipeline - "
        "from design and training through validation and integration into the Chrome extension "
        "- as well as the post-model intelligence layer (trusted domains, free email provider "
        "separation, newsletter detection), the AI Enhancement (BYOK) feature, BEC/linkless "
        "detection, attachment analysis, user-managed trusted domains, the risk scoring dashboard, "
        "and comprehensive documentation updates."
    )

    # -- ML Model Design --
    doc.add_heading("Machine Learning Model", level=2)
    doc.add_paragraph(
        "A unified phishing detection model will be designed and trained using Python and "
        "scikit-learn. The model will analyze 64 features extracted from emails, DNS records, "
        "page structure, BEC patterns, and attachment metadata. A single Random Forest "
        "classifier with isotonic calibration will produce a well-calibrated probability that "
        "maps directly to the 0-100 risk score. The training dataset will be augmented with "
        "synthetic BEC, attachment phishing, newsletter, and transactional email samples to "
        "improve accuracy on underrepresented phishing types."
    )

    doc.add_heading("Model Training", level=3)

    train_table = doc.add_table(rows=1, cols=2)
    train_table.style = "Table Grid"
    add_header_row(train_table, ["Artifact", "Description"])
    train_artifacts = [
        ("Training Script (train_model.py)",
         "Python script that will load the Kaggle dataset, define the 64-feature unified schema, "
         "generate synthetic training data (BEC, attachment, newsletter, transactional), augment "
         "the dataset with scan-scenario variants, train a Random Forest with CalibratedClassifierCV "
         "(isotonic, 5-fold CV), validate with classification reports and per-phishing-type "
         "evaluation, and export the full model to JSON."),
        ("Training Dataset",
         "Phishing_Legitimate_full.csv from Kaggle, augmented to ~36,000 samples with "
         "scan-scenario variants and synthetic phishing/legitimate samples covering BEC, "
         "attachment phishing, newsletters, and transactional emails."),
        ("Unified Model (model_unified.json)",
         "Single Random Forest with 200 decision trees trained on 64 features. Exported as JSON "
         "containing tree structures, StandardScaler parameters, and isotonic calibration lookup "
         "table for in-browser inference."),
        ("Feature Name File (feature_names_unified.json)",
         "Ordered list of 64 feature names expected by the unified model, matching "
         "buildUnifiedVector() in featureExtractor.js."),
        ("Calibrated Model (model_unified_calibrated.pkl)",
         "Scikit-learn CalibratedClassifierCV model for offline evaluation and retraining."),
    ]
    for a in train_artifacts:
        add_data_row(train_table, a)
    set_table_style(train_table)

    doc.add_heading("Target Model Performance", level=3)
    doc.add_paragraph(
        "The following metrics are targeted during model development."
    )

    perf_table = doc.add_table(rows=1, cols=2)
    perf_table.style = "Table Grid"
    add_header_row(perf_table, ["Metric", "Target"])
    perf_data = [
        ("Number of Features", "64 (unified schema across 7 groups)"),
        ("Feature Groups", "URL/Email (25), Custom Rules (9), DNS (5), Deep Scan (13), "
         "BEC/Linkless (5), Attachment (5), Context Flags (2)"),
        ("Model Type", "Random Forest + CalibratedClassifierCV (isotonic)"),
        ("Hyperparameters", "n_estimators=200, max_depth=20, min_samples_leaf=2, n_jobs=-1"),
        ("Calibration", "Isotonic regression via CalibratedClassifierCV (5-fold CV)"),
        ("Dataset Size", "~36,000 samples (augmented with synthetic data)"),
        ("Training/Test Split", "80% training / 20% test, stratified, random_state=42"),
        ("Target Accuracy", ">95% overall, 100% on synthetic BEC and attachment phishing"),
        ("Score Formula", "riskScore = round(100 x calibrated_probability)"),
        ("Validation Method", "Classification report, confusion matrix, feature importance, "
         "per-phishing-type evaluation"),
        ("Per-Type Evaluation", "URL-credential, linkless BEC, attachment-led, deep scan impersonation"),
    ]
    for p in perf_data:
        add_data_row(perf_table, p)
    set_table_style(perf_table)

    # -- Post-Model Intelligence --
    doc.add_heading("Post-Model Intelligence", level=2)
    doc.add_paragraph(
        "A post-model intelligence layer will be implemented to apply targeted score adjustments "
        "after ML inference, reducing false positives on legitimate email while ensuring phishing "
        "from free email providers is not incorrectly dampened."
    )

    pmi_bullets = [
        "Trusted domain whitelist: 500+ built-in corporate/organizational domains across 15+ "
        "categories (Big Tech, Finance, News, E-commerce, Government, Education, etc.)",
        "Free email provider separation: Domains like gmail.com, outlook.com, icloud.com, "
        "yahoo.com, protonmail.com will be explicitly excluded from trusted dampening since "
        "anyone can register accounts on these services",
        "BEC rule boosts: Emails with strong financial request, authority impersonation, or "
        "phone callback signals will have their score floor raised to 70-80",
        "Trusted domain dampening: Emails from trusted corporate domains with no BEC/attachment "
        "signals will be capped at a score of 30",
        "Newsletter detection: Emails with unsubscribe links, 'view in browser' links, and "
        "common newsletter footer text will be capped at 45",
        "User-managed custom domains: Users will be able to add/remove trusted domains or "
        "block built-in ones through the popup settings UI, stored in chrome.storage.local",
    ]
    for b in pmi_bullets:
        doc.add_paragraph(b, style="List Bullet")

    # -- Extension Integration --
    doc.add_heading("Extension Integration", level=2)
    doc.add_paragraph(
        "The trained model and supporting features will be integrated into the Chrome extension. "
        "The following source files will be created or modified during Sprint 2."
    )

    code_table = doc.add_table(rows=1, cols=2)
    code_table.style = "Table Grid"
    add_header_row(code_table, ["File", "Planned Changes"])
    code_files = [
        ("content.js",
         "Rewrite for unified model: loadModel(), predictWithCalibratedForest() with Z-score "
         "normalization + tree traversal + isotonic calibration, runInference() with 64-element "
         "vector and post-model intelligence (trusted domains, free email provider separation, "
         "BEC boosts, newsletter detection). Add AI Enhancement: buildAiPayload(), shouldCallAi() "
         "gating, runAiAnalysis(), displayAiResult() with agreement badge. Implement 500+ "
         "TRUSTED_DOMAINS set, FREE_EMAIL_PROVIDERS set, and user-managed domain support."),
        ("featureExtractor.js",
         "Extend FeatureExtractor with BEC/linkless detection methods and attachment analysis "
         "methods. Add buildUnifiedVector() to construct the 64-element feature array with "
         "default-filling for missing groups."),
        ("background.js",
         "Add 5 AI provider adapters with shared system prompt enforcing no-tools/no-browsing. "
         "Add validateAiResponse() for strict JSON schema validation. Add AI settings to default "
         "storage initialization."),
        ("popup.html",
         "Add 'Enhance with AI' toggle, AI configuration modal, and Trusted Domains Manager "
         "section (input field, Trust/Block buttons, chip-based list display)."),
        ("popup.js",
         "Add AI settings handling, modal logic, and trusted domains management (add/remove/"
         "block with chrome.storage.local persistence and live rendering)."),
        ("train_model.py",
         "Complete rewrite for unified model with synthetic data generation for BEC phishing, "
         "attachment phishing, legitimate newsletters, and transactional emails."),
        ("manifest.json",
         "Add host_permissions for AI provider APIs."),
    ]
    for c in code_files:
        add_data_row(code_table, c)
    set_table_style(code_table)

    # -- Risk Scoring Pipeline --
    doc.add_heading("Risk Scoring Pipeline", level=2)
    doc.add_paragraph(
        "The phishing risk score will be computed using a single unified calibrated model "
        "followed by post-model intelligence:"
    )
    doc.add_paragraph(
        "riskScore = round(100 x calibrated_probability)",
    )
    doc.add_paragraph(
        "After ML inference, post-model adjustments will apply: BEC rule boosts (floor 70-80), "
        "trusted domain dampening (cap 30), newsletter dampening (cap 45), with free email "
        "providers explicitly excluded from dampening. The final score maps to four risk levels:"
    )

    risk_bullets = [
        "Friendly Fish - Low risk emails (score 0-49)",
        "Suspicious Fish - Medium risk emails (score 50-75)",
        "Phishy Puffer - High risk emails (score 76-89)",
        "Mega Phish Shark - Dangerous emails (score 90-100)",
    ]
    for b in risk_bullets:
        doc.add_paragraph(b, style="List Bullet")

    # -- AI Enhancement --
    doc.add_heading("AI Enhancement (Cloud BYOK)", level=2)
    doc.add_paragraph(
        "An optional cloud AI enhancement will be implemented to provide a second opinion on "
        "email risk using the user's own API key (Bring Your Own Key). Key design principles:"
    )

    ai_bullets = [
        "Features-only payload: No email body, subject, or sender address will ever be sent to AI",
        "Automatic gating: AI only called when local model is uncertain (score 30-80, low "
        "confidence, risky signals)",
        "Multi-provider support: OpenAI, Anthropic, Google Gemini, Azure OpenAI, Custom endpoint",
        "Strict JSON schema: AI must return aiRiskScore, riskTier, phishType, topSignals, "
        "confidence, notes",
        "Prompt injection hardening: System prompt enforces no tools, no browsing, no link visiting",
        "Agreement badge: Shows 'Aligned' or 'Needs review' based on local vs AI score difference",
        "Keys stored locally: chrome.storage.local only, never synced",
    ]
    for b in ai_bullets:
        doc.add_paragraph(b, style="List Bullet")

    # -- Dashboard Features --
    doc.add_heading("Dashboard and UI Features", level=2)
    doc.add_paragraph(
        "The following user-facing features will be implemented during Sprint 2."
    )

    ui_table = doc.add_table(rows=1, cols=2)
    ui_table.style = "Table Grid"
    add_header_row(ui_table, ["Feature", "Description"])
    ui_features = [
        ("Risk Score Display",
         "Numeric risk score (0-100) from calibrated ML probability with post-model "
         "adjustments. Displayed in the side panel with color coding. Risk badge appears "
         "next to the email subject in Gmail."),
        ("Risk Reasons List",
         "Human-readable explanations derived from feature values and post-model adjustments. "
         "Includes explanations for trusted domain dampening, free email provider warnings, "
         "BEC boosts, and newsletter detection."),
        ("AI Analysis Section",
         "When AI Enhancement is enabled: 'AI: Scanning...' indicator, AI Risk Score with "
         "color coding, risk tier, top signals, and agreement badge."),
        ("AI Configuration Modal",
         "Provider selection, API key input, custom endpoint URL, model name override."),
        ("Trusted Domains Manager",
         "Settings section with domain input, Trust/Block buttons, chip-based list display "
         "with remove capability, live count showing built-in + custom counts."),
        ("Suspicious Links Panel",
         "All links analyzed individually. Suspicious links displayed with anchor text "
         "and actual destination."),
        ("Loading & Error States",
         "Spinner animation during initial scan, progress during Deep Scan, 'AI: Scanning...' "
         "during AI analysis, and clear error handling throughout."),
    ]
    for u in ui_features:
        add_data_row(ui_table, u)
    set_table_style(ui_table)

    # -- Documentation --
    doc.add_heading("Documentation", level=2)

    doc_table = doc.add_table(rows=1, cols=2)
    doc_table.style = "Table Grid"
    add_header_row(doc_table, ["Document", "Description"])
    docs = [
        ("ARCHITECTURE.md",
         "Rewrite with updated Mermaid diagrams: system architecture with AI providers and "
         "trusted domains, risk scoring pipeline with post-model intelligence layer, data "
         "storage schema with custom domain fields, and email scan sequence with free email "
         "provider awareness."),
        ("README.md",
         "Update features table, add trusted domains section, add post-model intelligence "
         "description, update model accuracy, add free email provider explanation, update "
         "testing section."),
        ("Framework/ Diagrams",
         "Regenerate all 6 architecture diagram images and architecture PDF to reflect "
         "trusted domains, post-model intelligence, free email provider separation, and "
         "updated model accuracy (96.6%)."),
        ("FEATURE_MAPPING.md",
         "Document all 64 features across 7 groups with descriptions, sources, and positions."),
        ("Sprint 2 Documents",
         "Sprint 2 Requirements and Artifacts documents (this file)."),
    ]
    for d in docs:
        add_data_row(doc_table, d)
    set_table_style(doc_table)

    # -- Summary --
    doc.add_heading("Sprint 2 Deliverables Summary", level=2)

    summary_table = doc.add_table(rows=1, cols=2)
    summary_table.style = "Table Grid"
    add_header_row(summary_table, ["Metric", "Target"])
    summary_data = [
        ("Requirements to Complete", "10 (+ 10 stretch goals)"),
        ("Story Points", "49 (+ stretch)"),
        ("Source Files to Create/Modify", "7 (train_model.py, content.js, featureExtractor.js, "
         "background.js, popup.html, popup.js, manifest.json)"),
        ("ML Features", "64 (unified schema, 7 groups)"),
        ("Decision Trees", "200 (with isotonic calibration)"),
        ("Target Model Accuracy", ">95% overall, 100% on synthetic BEC/attachment phishing"),
        ("Trusted Domains", "500+ built-in + user custom"),
        ("Free Email Providers", "40+ explicitly excluded from dampening"),
        ("AI Providers to Support", "5 (OpenAI, Anthropic, Google, Azure, Custom)"),
        ("Risk Score Range", "0 - 100 (4 risk levels)"),
        ("Documentation Updates", "README, ARCHITECTURE, FEATURE_MAPPING, Framework/, Sprint 2 docs"),
        ("Team Members", "5"),
    ]
    for s in summary_data:
        add_data_row(summary_table, s)
    set_table_style(summary_table)

    out = os.path.join(OUT_DIR, "Sprint2_Artifacts.docx")
    doc.save(out)
    print(f"  -> {out}")


# ===================================================================
def _clean_pycache():
    """Remove __pycache__ from project root (Chrome rejects dirs starting with _)."""
    import shutil
    pc = os.path.join(BASE_DIR, '__pycache__')
    if os.path.isdir(pc):
        shutil.rmtree(pc, ignore_errors=True)
        print("  (cleaned __pycache__)")


if __name__ == "__main__":
    print("Generating Sprint 2 documents...")
    build_requirements()
    build_artifacts()
    _clean_pycache()
    print("Done!")
